from src.interfaces.IParser import IParser
from typing import Optional, Iterator
import docx
from src.interfaces.parser import ParserInfo, FileType, Connectivity


class DocxParser(IParser):
    info = ParserInfo(
        name= "docx_default",
        supported_types= [FileType.DOCX],
        connectivity=Connectivity.OFFLINE,
        is_ocr=False,
    )

    doc = None
    block_iterator: Optional[Iterator[str]] = None

    def __init__(self) -> None:
        pass
    def open(self, file_name: Optional[str]) -> None:
        try:
            if not file_name:
                raise ValueError("File name is required")

            if not file_name.endswith(".docx"):
                raise ValueError("DocxParser only supports .docx files")

            self.file_name = file_name
            self.doc = docx.Document(self.file_name)

            self.block_iterator = self._extract_text_blocks()

        except Exception as e:
            raise Exception(f"Error opening docx file: {e}")

    def _extract_text_blocks(self) -> Iterator[str]:
        for para in self.doc.paragraphs:
            text = para.text.strip()
            if text:
                yield text

        for table in self.doc.tables:
            for row in table.rows:
                row_text = [cell.text.strip() for cell in row.cells]
                yield " | ".join(row_text)

    def get_next_text_block(self) -> Optional[str]:
        try:
            if self.block_iterator is None:
                raise Exception("File is not open")

            return next(self.block_iterator)

        except StopIteration:
            return None
        except Exception as e:
            print(f"Error reading docx block: {e}")
            return None

    def close(self) -> None:
        self.doc = None
        self.block_iterator = None